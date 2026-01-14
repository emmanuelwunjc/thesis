#!/usr/bin/env python3
"""
Add in-text references properly - only in main body sections.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING
import os

def add_references(input_path, output_path):
    print(f"📄 Loading: {input_path}")
    doc = Document(input_path)
    
    # Track which section we're in
    current_section = None
    sections_order = []
    
    # First pass: identify sections
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text in ["Abstract", "Introduction and Background", "Literature Review", 
                    "Conceptual Framework", "Data and Methods", "Results", 
                    "Discussion", "Contributions", "Limitations", "References", "APPENDIX"]:
            sections_order.append((i, text))
    
    print(f"   Found sections: {[s[1] for s in sections_order]}")
    
    # Build section ranges
    section_ranges = {}
    for idx, (para_idx, section_name) in enumerate(sections_order):
        end_idx = sections_order[idx + 1][0] if idx + 1 < len(sections_order) else len(doc.paragraphs)
        section_ranges[section_name] = (para_idx, end_idx)
    
    refs_added = []
    
    # Second pass: add references in appropriate sections
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Skip headings, captions, empty
        if not text or text in section_ranges or text.startswith("Table ") or text.startswith("Figure "):
            continue
        
        # Determine current section
        current_section = None
        for section, (start, end) in section_ranges.items():
            if start <= i < end:
                current_section = section
                break
        
        # Skip Abstract and References sections
        if current_section in ["Abstract", "References", "APPENDIX", None]:
            continue
        
        # === CONCEPTUAL FRAMEWORK: Add Figure 1 reference ===
        if current_section == "Conceptual Framework":
            if "These relationships lead to four hypotheses" in text:
                new_text = text.replace(
                    "These relationships lead to four hypotheses",
                    "Figure 1 illustrates these key relationships and hypothesized pathways. These relationships lead to four hypotheses"
                )
                para.clear()
                run = para.add_run(new_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                refs_added.append("Figure 1 in Conceptual Framework")
                print(f"   ✅ Added Figure 1 ref at para {i}")
        
        # === DATA AND METHODS: Add Table 1 reference ===
        if current_section == "Data and Methods":
            if "analytic sample" in text.lower() and "2,421" in text:
                if "(Table 1" not in text:
                    new_text = text.replace("2,421", "2,421 (see Table 1 for descriptive statistics)")
                    para.clear()
                    run = para.add_run(new_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    refs_added.append("Table 1 in Data and Methods")
                    print(f"   ✅ Added Table 1 ref at para {i}")
            
            # Add Appendix E reference for robustness
            if "robustness" in text.lower() and "Appendix E" not in text:
                new_text = text + " (see Appendix E for alternative specifications)"
                para.clear()
                run = para.add_run(new_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                refs_added.append("Appendix E in Data and Methods")
                print(f"   ✅ Added Appendix E ref at para {i}")
        
        # === RESULTS: Add Table 2 and Table 3 references ===
        if current_section == "Results":
            # Main model results - Table 2
            if "privacy caution" in text.lower() and ("strongest" in text.lower() or "−2.3" in text or "-2.3" in text):
                if "(Table 2" not in text and "(Table 3" not in text:
                    new_text = text + " (Table 2)"
                    para.clear()
                    run = para.add_run(new_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    refs_added.append("Table 2 in Results")
                    print(f"   ✅ Added Table 2 ref at para {i}")
            
            # Interaction model - Table 3
            if "interaction" in text.lower() and ("significant" in text.lower() or "0.49" in text):
                if "diabetes" in text.lower() and "(Table 3" not in text:
                    new_text = text + " (Table 3)"
                    para.clear()
                    run = para.add_run(new_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    refs_added.append("Table 3 in Results")
                    print(f"   ✅ Added Table 3 ref at para {i}")
        
        # === DISCUSSION: Add Appendix C reference for age findings ===
        if current_section == "Discussion":
            if "age" in text.lower() and ("older" in text.lower() or "younger" in text.lower()):
                if "Appendix C" not in text and "subgroup" not in text.lower():
                    # Only add once
                    if "Appendix C in Discussion" not in refs_added:
                        new_text = text + " (see Appendix C for age subgroup analysis)"
                        para.clear()
                        run = para.add_run(new_text)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        refs_added.append("Appendix C in Discussion")
                        print(f"   ✅ Added Appendix C ref at para {i}")
    
    print(f"   📊 Total references added: {len(refs_added)}")
    
    print(f"💾 Saving: {output_path}")
    doc.save(output_path)
    print("✅ Done!")

def main():
    # Start fresh from the version with images embedded
    input_dir = "/Users/wuyiming/code/thesis/submission"
    
    add_references(
        os.path.join(input_dir, "Wu_DataSharing_SUBMIT.docx"),
        os.path.join(input_dir, "Wu_DataSharing_FINAL_SUBMIT.docx")
    )

if __name__ == "__main__":
    main()
