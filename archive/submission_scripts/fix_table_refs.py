#!/usr/bin/env python3
"""
Fix table references in Results section - ensure Table 2 mentioned first.
"""

from docx import Document
from docx.shared import Pt
import os

def fix_references(input_path, output_path):
    print(f"📄 Loading: {input_path}")
    doc = Document(input_path)
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        
        # Find the Results paragraph and fix the table references
        if "Privacy caution is one of the major factors" in text:
            # This is the first paragraph of Results
            # Add proper references
            new_text = text.replace(
                "Table 3 presents the interaction model results.",
                "Table 2 presents the main regression results; Table 3 presents the interaction model."
            )
            # Also fix if Table 2 wasn't added properly
            if "Table 2 presents" not in new_text:
                new_text = new_text.replace(
                    "In general, the findings",
                    "Table 2 presents the main model results. In general, the findings"
                )
            
            para.clear()
            run = para.add_run(new_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            print(f"   ✅ Fixed table references at para {i}")
            break
    
    print(f"💾 Saving: {output_path}")
    doc.save(output_path)
    print("✅ Done!")

def main():
    input_dir = "/Users/wuyiming/code/thesis/submission"
    
    fix_references(
        os.path.join(input_dir, "Wu_DataSharing_FINAL_SUBMIT.docx"),
        os.path.join(input_dir, "Wu_DataSharing_FINAL_SUBMIT.docx")  # Overwrite
    )

if __name__ == "__main__":
    main()
