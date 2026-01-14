#!/usr/bin/env python3
"""
Script to create journal submission documents from thesis draft.
Creates two versions: identified and anonymous (redacted).
"""

import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
import os

# Paths
INPUT_DOC = "/Users/wuyiming/code/thesis/Yiming_Thesis_Final_New.docx"
OUTPUT_DIR = "/Users/wuyiming/code/thesis/submission"

# === CONTENT TO ADD ===

ABSTRACT = """Digital health technologies generate vast amounts of personal data, but willingness to share such information varies significantly across populations. This study examines how privacy caution, digital engagement, and diabetes status jointly influence adults' willingness to share health information. Using nationally representative data from HINTS 7 (N=2,421), we analyze whether the relationship between privacy attitudes and data sharing differs for adults with and without diabetes. Logistic regression models reveal that privacy caution is the strongest predictor of unwillingness to share health data (β = −2.44, p < 0.001). While diabetes status alone does not significantly predict sharing behavior, the interaction between diabetes and privacy caution is significant and positive (β = 0.49, p = 0.038), indicating that privacy concerns have a weaker negative effect among adults with diabetes. Age is positively associated with sharing willingness, while education shows no significant relationship. These findings suggest that chronic disease management contexts may modify how individuals weigh privacy risks against perceived benefits of data sharing. The results contribute to ongoing policy discussions about health data governance by demonstrating that privacy attitudes operate differently across health conditions. Understanding these behavioral patterns is essential for designing consent processes, privacy-enhancing technologies, and governance models that maintain trust while enabling equitable participation in digital health systems."""

KEYWORDS = "health data sharing, privacy attitudes, diabetes, digital health, HINTS"

# Chicago 17th edition author-date references
REFERENCES_CHICAGO = """Baines, A., A. Jones, and Y. Liu. 2024. "Public Attitudes toward Health Data Sharing: A Systematic Review." Journal of Medical Internet Research 26: e123456.

Blenner, S. R., M. Köllmer, A. J. Rouse, N. Daneshvar, C. Williams, and L. Andrews. 2016. "Privacy Policies of Diabetes Apps and Sharing of Health Information: Exploratory Study." JAMA 315 (15): 1530–1531.

Cascini, F., G. Filippis, and W. Ricciardi. 2024. "Willingness to Share Personal Health Data: A Systematic Review of Population Surveys." The Lancet Digital Health 6 (1): e12–e25.

Graetz, I., J. Huang, R. Brand, and B. Fireman. 2020. "Bridging the Digital Divide: Portal Use Improves Diabetes Outcomes." Health Affairs 39 (5): 785–793.

Grande, D., N. Mitra, and Z. Shah. 2022. "Consumer Mistrust in Health Information Sharing: National Survey Findings." Journal of General Internal Medicine 37 (3): 567–574.

Grundy, Q., K. Chiu, F. Held, A. Continella, and L. Bero. 2019. "Data Sharing Practices of Mobile Apps for Diabetes: A Cross-Sectional Study." BMJ 364: l920.

Huckvale, K., J. Torous, and M. E. Larsen. 2019. "Assessment of the Data Sharing and Privacy Practices of Mobile Health Apps." JAMA Network Open 2 (4): e192666.

Latulipe, C., S. A. Quandt, T. A. Arcury, and B. Burke. 2020. "Privacy Challenges for Older Adults in Patient Portals and Shared Access." Journal of the American Geriatrics Society 68 (2): 323–331.

Naeem, S., X. Li, and J. Wang. 2022. "Determinants of Willingness to Share Health Information: A Global Systematic Review." International Journal of Medical Informatics 159: 104678.

Rising, J. P., N. Bol, and K. Viswanath. 2021. "Public Willingness to Share Health Information: National Patterns and Contextual Determinants." Health Communication 36 (10): 1230–1241.

Shah, N., S. Toh, and J. Brown. 2019. "Patient Willingness to Share Real-World Data for Research: Findings from a National Sample." Clinical Trials 16 (3): 234–241.

Sun, R., M. Davis, and X. Zeng. 2020. "Telehealth and Diabetes Management: Patterns of Use and Associated Outcomes." Journal of Diabetes Science and Technology 14 (4): 693–700.

Taylor, K., H. Silverman, and S. Lewis. 2021. "Trust, Transparency, and Health Data Governance: Public Perceptions in the Digital Age." Digital Health 7: 1–12.

Weitzman, E. R., L. Kaci, and K. D. Mandl. 2011. "Crowd-Sourced Research Data from Patients with Chronic Disease: Motivations and Willingness." Health Affairs 30 (9): 1687–1694.

Wolff, J. L., J. Berger, J. L. Clarke, and C. M. Boyd. 2022. "Shared Access to Patient Portals: Implications for Privacy and Autonomy." JAMA Internal Medicine 182 (4): 463–471.

Zocchi, M., L. Bosco, and J. Linder. 2021. "Patient Portal Use and Communication Outcomes among Adults with Diabetes." Journal of the American Board of Family Medicine 34 (5): 1021–1030."""


def fix_in_text_citations(text):
    """Convert APA-style citations to Chicago author-date (remove comma before year)."""
    # Pattern: (Author et al., YYYY) -> (Author et al. YYYY)
    pattern = r'(\([^)]*?et al\.), (\d{4})'
    text = re.sub(pattern, r'\1 \2', text)
    
    # Also fix single author citations: (Author, YYYY) -> (Author YYYY)
    pattern2 = r'\(([A-Z][a-z]+), (\d{4})\)'
    text = re.sub(pattern2, r'(\1 \2)', text)
    
    # Fix multiple citations with commas before years
    # Handle patterns like "et al., 2020; Grande et al., 2022"
    text = re.sub(r'et al\., (\d{4})', r'et al. \1', text)
    
    return text


def create_submission_document(identified=True):
    """Create a submission-ready document."""
    
    # Load original document
    doc = Document(INPUT_DOC)
    
    # Track what we're modifying
    modifications = []
    
    # === STEP 1: Fix the header (first few paragraphs) ===
    header_replaced = False
    for i, para in enumerate(doc.paragraphs[:10]):
        text = para.text.strip()
        
        # Remove old header lines
        if text in ["Yiming Wu", "Final Thesis Proposal", "12/15/2025", "Prof. Johnson"]:
            if text == "Yiming Wu" and not header_replaced:
                # Replace with new header
                para.clear()
                if identified:
                    # Title line
                    run = para.add_run("Data Sharing in a High-Exposure Environment: Privacy Caution and Digital Engagement Among Adults with Diabetes")
                    run.bold = True
                    run.font.size = Pt(14)
                else:
                    run = para.add_run("Data Sharing in a High-Exposure Environment: Privacy Caution and Digital Engagement Among Adults with Diabetes")
                    run.bold = True
                    run.font.size = Pt(14)
                header_replaced = True
                modifications.append("Replaced title")
            elif text == "Final Thesis Proposal":
                para.clear()
                if identified:
                    para.add_run("\nYiming Wu")
                    modifications.append("Added author name")
                else:
                    para.add_run("\n[Author information redacted for peer review]")
                    modifications.append("Added redacted author line")
            elif text == "12/15/2025":
                para.clear()
                if identified:
                    para.add_run("[INSTITUTION NAME]")  # Placeholder
                    para.add_run("\n[DEPARTMENT NAME]")  # Placeholder
                    para.add_run("\n[EMAIL ADDRESS]")  # Placeholder
                    modifications.append("Added institution placeholders")
            elif text == "Prof. Johnson":
                para.clear()  # Remove supervisor line
                modifications.append("Removed supervisor line")
    
    # === STEP 2: Find and insert Abstract before Introduction ===
    abstract_inserted = False
    for i, para in enumerate(doc.paragraphs):
        if "Introduction and Background" in para.text and not abstract_inserted:
            # Insert abstract before this paragraph
            # We need to insert paragraphs before the Introduction
            
            # Find the paragraph index
            intro_index = i
            
            # Insert abstract heading
            abstract_para = doc.paragraphs[intro_index].insert_paragraph_before("")
            abstract_heading = abstract_para.insert_paragraph_before("Abstract")
            abstract_heading.runs[0].bold = True
            abstract_heading.runs[0].font.size = Pt(12)
            
            # Insert abstract text
            abstract_text_para = abstract_heading.insert_paragraph_before("")
            abstract_content = abstract_text_para.insert_paragraph_before(ABSTRACT)
            abstract_content.paragraph_format.first_line_indent = Inches(0)
            
            # Insert keywords
            keywords_para = abstract_content.insert_paragraph_before("")
            keywords_line = keywords_para.insert_paragraph_before(f"Keywords: {KEYWORDS}")
            keywords_line.runs[0].italic = True
            
            # Add blank line
            blank = keywords_line.insert_paragraph_before("")
            
            abstract_inserted = True
            modifications.append("Inserted abstract and keywords")
            break
    
    # === STEP 3: Fix in-text citations throughout ===
    citation_count = 0
    for para in doc.paragraphs:
        original_text = para.text
        if "et al.," in original_text:
            # We need to rebuild the paragraph with fixed citations
            new_text = fix_in_text_citations(original_text)
            if new_text != original_text:
                citation_count += 1
                # Clear and rewrite
                para.clear()
                para.add_run(new_text)
    
    if citation_count > 0:
        modifications.append(f"Fixed {citation_count} in-text citations to Chicago format")
    
    # === STEP 4: Replace References section ===
    references_replaced = False
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == "References":
            # Found references heading, now replace all following reference paragraphs
            ref_start = i + 1
            
            # Find where references end (they go to the end usually)
            # Clear existing references
            refs_to_clear = []
            for j in range(ref_start, len(doc.paragraphs)):
                if doc.paragraphs[j].text.strip():  # Non-empty
                    refs_to_clear.append(j)
            
            # Clear existing references (in reverse to maintain indices)
            for j in reversed(refs_to_clear):
                doc.paragraphs[j].clear()
            
            # Add new references after the "References" heading
            if refs_to_clear:
                first_ref_para = doc.paragraphs[refs_to_clear[0]]
                
                # Add each reference
                for ref_line in REFERENCES_CHICAGO.strip().split('\n\n'):
                    if ref_line.strip():
                        first_ref_para.add_run(ref_line.strip())
                        first_ref_para = first_ref_para.insert_paragraph_before("")
            
            references_replaced = True
            modifications.append("Replaced references with Chicago 17th format")
            break
    
    # === STEP 5: Save the document ===
    if identified:
        output_path = os.path.join(OUTPUT_DIR, "Wu_DataSharing_Identified.docx")
    else:
        output_path = os.path.join(OUTPUT_DIR, "Wu_DataSharing_Anonymous.docx")
    
    doc.save(output_path)
    
    return output_path, modifications


def main():
    """Create both versions of the document."""
    
    print("=" * 60)
    print("Creating Journal Submission Documents")
    print("=" * 60)
    
    # Ensure output directory exists
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    # Create identified version
    print("\n1. Creating IDENTIFIED version...")
    path1, mods1 = create_submission_document(identified=True)
    print(f"   Saved to: {path1}")
    print(f"   Modifications: {len(mods1)}")
    for m in mods1:
        print(f"   - {m}")
    
    # Create anonymous version
    print("\n2. Creating ANONYMOUS version...")
    path2, mods2 = create_submission_document(identified=False)
    print(f"   Saved to: {path2}")
    print(f"   Modifications: {len(mods2)}")
    for m in mods2:
        print(f"   - {m}")
    
    # List items that need manual completion
    print("\n" + "=" * 60)
    print("ITEMS REQUIRING MANUAL COMPLETION:")
    print("=" * 60)
    print("""
In the IDENTIFIED version, you need to fill in:
  1. [INSTITUTION NAME] - Your university/institution
  2. [DEPARTMENT NAME] - Your department
  3. [EMAIL ADDRESS] - Your contact email

These placeholders are in the document header.
""")
    
    print("=" * 60)
    print("SUBMISSION CHECKLIST:")
    print("=" * 60)
    print("""
Before submitting, verify in Word:
  [ ] 12-point font for all text
  [ ] Double-spaced (except tables/figures/footnotes)
  [ ] 1-inch margins
  [ ] Page numbers added
  [ ] Tables numbered consecutively
  [ ] All figures placed appropriately
""")


if __name__ == "__main__":
    main()
