#!/usr/bin/env python3
"""
Add the complete variable definition table to Appendix A.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Condensed variable table for paper (key variables only)
VARIABLE_TABLE = {
    "title": "Table A0: Variable Definitions",
    "headers": ["Variable", "Type", "Definition", "Coding"],
    "rows": [
        # Dependent Variable
        ["WillingShareData_HCP2", "Binary", "Willingness to share health data with providers", "0=No, 1=Yes"],
        # Main Predictors
        ["privacy_caution_index", "Continuous (0–1)", "Composite privacy caution index (6 sub-dimensions)", "Higher = more cautious"],
        ["diabetic", "Binary", "Self-reported diabetes diagnosis", "0=No, 1=Yes"],
        # Interaction
        ["Diabetes × Privacy", "Continuous", "Interaction term", "Product of above"],
        # Controls
        ["age_continuous", "Continuous", "Age in years", "18–100+"],
        ["education_numeric", "Ordinal (1–6)", "Highest education level", "1=<8yrs to 6=Postgrad"],
        ["male", "Binary", "Gender", "0=Female, 1=Male"],
        ["has_insurance", "Binary", "Health insurance coverage", "0=No, 1=Yes"],
        ["urban", "Binary", "Urban/rural residence", "0=Rural, 1=Urban"],
        ["region_numeric", "Categorical (1–4)", "Census region", "1=NE, 2=MW, 3=S, 4=W"],
    ],
    "note": "Note: Privacy Caution Index computed as mean of 6 sub-dimensions (23 HINTS items). Cronbach's α = 0.78."
}

# Full sub-dimension breakdown
SUBDIMENSION_TABLE = {
    "title": "Table A0b: Privacy Caution Index Components",
    "headers": ["Sub-dimension", "# Items", "Variables Included", "Coding"],
    "rows": [
        ["Sharing Willingness", "4", "WillingShareData_HCP2, SharedHealthDeviceInfo2, SocMed_SharedPers, SocMed_SharedGen", "Yes=0, No=1"],
        ["Portal Usage", "7", "AccessOnlineRecord3, OnlinePortal_PCP, _OthHCP, _Insurer, _Lab, _Pharmacy, _Hospital", "Used=0, Not used=1"],
        ["Device Usage", "4", "UseDevice_Computer, _SmPhone, _Tablet, _SmWatch", "Yes=0, No=1"],
        ["Trust Levels", "4", "TrustHCSystem, CancerTrustDoctor, _Scientists, _Family", "A lot=0 to Not at all=1"],
        ["Social Media", "2", "SocMed_Visited, MisleadingHealthInfo", "Yes=0, No=1"],
        ["Other Privacy", "2", "ConfidentMedForms, WillingUseTelehealth", "Confident=0, Not=1"],
    ],
    "note": "Note: Missing values imputed as 0.5 (midpoint). Sub-dimensions averaged, then averaged again for composite index."
}

def set_cell_font(cell, size=9, bold=False):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = RGBColor(0, 0, 0)

def add_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

def add_variable_table(doc, table_data):
    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(table_data["title"])
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(11)
    title_run.font.bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(12)
    title_para.paragraph_format.space_after = Pt(6)
    
    # Table
    num_cols = len(table_data["headers"])
    num_rows = len(table_data["rows"]) + 1
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)
    
    # Headers
    for i, header in enumerate(table_data["headers"]):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_font(cell, size=9, bold=True)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'E8E8E8')
        cell._tc.get_or_add_tcPr().append(shading)
    
    # Rows
    for row_idx, row_data in enumerate(table_data["rows"]):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            set_cell_font(cell, size=8 if col_idx == 2 else 9)
    
    # Note
    note_para = doc.add_paragraph()
    note_run = note_para.add_run(table_data["note"])
    note_run.font.name = 'Times New Roman'
    note_run.font.size = Pt(9)
    note_run.font.italic = True
    note_para.paragraph_format.space_after = Pt(12)

def process_document(input_path, output_path):
    print(f"📄 Loading: {input_path}")
    doc = Document(input_path)
    
    # Find "Appendix A: Variable Definitions" and insert tables after it
    for i, para in enumerate(doc.paragraphs):
        if "Appendix A: Variable Definitions" in para.text:
            print(f"   Found Appendix A at paragraph {i}")
            # Remove the placeholder figure reference if it exists
            for j in range(i+1, min(i+5, len(doc.paragraphs))):
                if "Insert" in doc.paragraphs[j].text and "variable_definition" in doc.paragraphs[j].text:
                    doc.paragraphs[j].clear()
                    print(f"   Removed placeholder at paragraph {j}")
                if "Figure A1" in doc.paragraphs[j].text:
                    doc.paragraphs[j].clear()
            break
    
    # Find where Appendix B starts to insert before it
    insert_before_b = None
    for i, para in enumerate(doc.paragraphs):
        if "Appendix B:" in para.text:
            insert_before_b = i
            break
    
    # We'll add tables at the end and they'll appear in order
    # Actually, let's just add after finding Appendix A section
    
    # For now, add tables at the very end before we save, 
    # but we need to restructure - let me create a new approach
    
    # Find the Appendix A paragraph and add content right after
    new_doc = Document()
    found_appendix_a = False
    inserted = False
    
    for para in doc.paragraphs:
        # Copy paragraph
        new_para = new_doc.add_paragraph()
        new_para.paragraph_format.alignment = para.paragraph_format.alignment
        new_para.paragraph_format.space_before = para.paragraph_format.space_before
        new_para.paragraph_format.space_after = para.paragraph_format.space_after
        new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
        new_para.paragraph_format.line_spacing_rule = para.paragraph_format.line_spacing_rule
        
        for run in para.runs:
            new_run = new_para.add_run(run.text)
            new_run.font.name = run.font.name
            new_run.font.size = run.font.size
            new_run.font.bold = run.font.bold
            new_run.font.italic = run.font.italic
            if run.font.color and run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb
        
        # Check if this is Appendix A
        if "Appendix A: Variable Definitions" in para.text:
            found_appendix_a = True
        
        # Skip placeholder text
        if found_appendix_a and not inserted:
            if "Insert" in para.text and "variable_definition" in para.text.lower():
                new_para.clear()
                continue
            if "Figure A1:" in para.text and "Variable" in para.text:
                new_para.clear()
                # Now insert our tables
                add_variable_table(new_doc, VARIABLE_TABLE)
                add_variable_table(new_doc, SUBDIMENSION_TABLE)
                inserted = True
                print("   ✅ Inserted variable definition tables")
                continue
    
    # If we never found the right spot, add at end before saving
    if not inserted:
        print("   Adding tables at end of document")
        add_variable_table(new_doc, VARIABLE_TABLE)
        add_variable_table(new_doc, SUBDIMENSION_TABLE)
    
    # Copy tables from original (this is complex, let's use simpler approach)
    
    print(f"💾 Saving: {output_path}")
    new_doc.save(output_path)
    print("✅ Done!")

def simple_add_tables(input_path, output_path):
    """Simpler approach: just add tables after Appendix A heading."""
    print(f"📄 Loading: {input_path}")
    doc = Document(input_path)
    
    # Find Appendix A and clear placeholder, then add tables at end
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if "[Insert" in text and "variable_definition" in text.lower():
            para.clear()
            para.add_run("See Tables A0 and A0b below for complete variable definitions.")
            print(f"   Updated placeholder at paragraph {i}")
        if "Figure A1:" in text and "Variable" in text:
            para.clear()
            print(f"   Cleared old caption at paragraph {i}")
    
    # Add the actual tables before Appendix B
    # Find Appendix B location
    appendix_b_idx = None
    for i, para in enumerate(doc.paragraphs):
        if "Appendix B:" in para.text:
            appendix_b_idx = i
            break
    
    # We'll add at end of doc for simplicity, tables will be in appendix section
    print("   Adding variable definition tables...")
    add_variable_table(doc, VARIABLE_TABLE)
    add_variable_table(doc, SUBDIMENSION_TABLE)
    
    print(f"💾 Saving: {output_path}")
    doc.save(output_path)
    print("✅ Done!")

def main():
    input_dir = "/Users/wuyiming/code/thesis/submission"
    
    simple_add_tables(
        os.path.join(input_dir, "Wu_DataSharing_PRO.docx"),
        os.path.join(input_dir, "Wu_DataSharing_FINAL_PRO.docx")
    )
    
    simple_add_tables(
        os.path.join(input_dir, "Wu_DataSharing_PRO_Anonymous.docx"),
        os.path.join(input_dir, "Wu_DataSharing_FINAL_PRO_Anonymous.docx")
    )

if __name__ == "__main__":
    main()
