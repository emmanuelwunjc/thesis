#!/usr/bin/env python3
"""
Add in-text references to figures, tables, and appendices in the main body.
"""

from docx import Document
from docx.shared import Pt
import os
import re

def add_intext_references(input_path, output_path):
    print(f"📄 Loading: {input_path}")
    doc = Document(input_path)
    
    refs_added = {
        "Figure 1": False,
        "Table 1": False,
        "Table 2": False,
        "Table 3": False,
        "Appendix C": False,
        "Appendix D": False,
        "Appendix E": False,
    }
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        
        # Skip if already has references or is a heading/caption
        if "(see " in text.lower() or "(Figure" in text or "(Table" in text:
            continue
        if text.startswith("Table ") or text.startswith("Figure ") or text.startswith("Appendix"):
            continue
        
        # === FIGURE 1: Conceptual Framework ===
        # Add after "conceptual framework" or "conceptual model" mention
        if not refs_added["Figure 1"]:
            if "conceptual" in text.lower() and ("framework" in text.lower() or "model" in text.lower()):
                if "These relationships lead to four hypotheses" in text:
                    # Add reference before the hypotheses
                    new_text = text.replace(
                        "These relationships lead to four hypotheses",
                        "Figure 1 illustrates these relationships. These relationships lead to four hypotheses"
                    )
                    para.clear()
                    run = para.add_run(new_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    refs_added["Figure 1"] = True
                    print(f"   ✅ Added Figure 1 reference at para {i}")
        
        # === TABLE 1: Descriptive Statistics ===
        # Add in Data and Methods when describing the sample
        if not refs_added["Table 1"]:
            if "2,421" in text and ("sample" in text.lower() or "observation" in text.lower()):
                if "analytic sample" in text.lower() or "final" in text.lower():
                    # Find good insertion point
                    if not "(see" in text:
                        new_text = text.replace("2,421", "2,421 (Table 1 presents descriptive statistics)")
                        para.clear()
                        run = para.add_run(new_text)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        refs_added["Table 1"] = True
                        print(f"   ✅ Added Table 1 reference at para {i}")
        
        # === TABLE 2 & 3: Regression Results ===
        # Add in Results section when discussing regression findings
        if not refs_added["Table 2"]:
            if "privacy caution" in text.lower() and "−2.3" in text or "-2.3" in text:
                if "coefficient" in text.lower() or "β" in text or "effect" in text.lower():
                    new_text = text + " (Table 2)"
                    para.clear()
                    run = para.add_run(new_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    refs_added["Table 2"] = True
                    print(f"   ✅ Added Table 2 reference at para {i}")
        
        if not refs_added["Table 3"]:
            if "interaction" in text.lower() and ("0.49" in text or "0.4896" in text or "significant" in text.lower()):
                if "diabetes" in text.lower() and "privacy" in text.lower():
                    new_text = text + " (Table 3)"
                    para.clear()
                    run = para.add_run(new_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    refs_added["Table 3"] = True
                    print(f"   ✅ Added Table 3 reference at para {i}")
        
        # === APPENDIX C: Age subgroup ===
        if not refs_added["Appendix C"]:
            if "age" in text.lower() and ("subgroup" in text.lower() or "heterogen" in text.lower()):
                if "older" in text.lower() or "younger" in text.lower() or "65" in text:
                    new_text = text + " (see Appendix C for age subgroup analysis)"
                    para.clear()
                    run = para.add_run(new_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    refs_added["Appendix C"] = True
                    print(f"   ✅ Added Appendix C reference at para {i}")
        
        # === APPENDIX E: Robustness checks ===
        if not refs_added["Appendix E"]:
            if "robust" in text.lower() or ("alternative" in text.lower() and "method" in text.lower()):
                new_text = text + " (see Appendix E for robustness checks)"
                para.clear()
                run = para.add_run(new_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                refs_added["Appendix E"] = True
                print(f"   ✅ Added Appendix E reference at para {i}")
    
    print(f"   📊 References added: {sum(refs_added.values())}/{len(refs_added)}")
    
    # If some refs weren't added automatically, let's add them at key sections
    # Find sections and add references
    section_refs = {
        "Conceptual Framework": ("Figure 1 presents the conceptual framework guiding this analysis.", "Figure 1"),
        "Data and Methods": ("Table 1 presents descriptive statistics for the analytic sample.", "Table 1"),
        "Results": ("Tables 2 and 3 present the main regression results.", "Table 2"),
    }
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Check if this is a section heading
        for section, (ref_text, ref_key) in section_refs.items():
            if text == section and not refs_added.get(ref_key, True):
                # Add reference sentence after the heading
                if i + 1 < len(doc.paragraphs):
                    next_para = doc.paragraphs[i + 1]
                    if next_para.text.strip() and not next_para.text.startswith("Table"):
                        # Prepend reference to next paragraph
                        original = next_para.text
                        next_para.clear()
                        run = next_para.add_run(ref_text + " " + original)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        refs_added[ref_key] = True
                        print(f"   ✅ Added {ref_key} reference after {section} heading")
    
    print(f"💾 Saving: {output_path}")
    doc.save(output_path)
    print("✅ Done!")

def main():
    input_dir = "/Users/wuyiming/code/thesis/submission"
    
    add_intext_references(
        os.path.join(input_dir, "Wu_DataSharing_SUBMIT.docx"),
        os.path.join(input_dir, "Wu_DataSharing_FINAL_SUBMIT.docx")
    )

if __name__ == "__main__":
    main()
