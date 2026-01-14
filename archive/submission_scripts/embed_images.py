#!/usr/bin/env python3
"""
Embed actual images into the document, replacing placeholders.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def embed_images(input_path, output_path):
    print(f"📄 Loading: {input_path}")
    doc = Document(input_path)
    
    figures_dir = "/Users/wuyiming/code/thesis/figures"
    
    # Images to embed
    images = {
        "Figure 1": os.path.join(figures_dir, "conceptual_framework_diagram.png"),
        "Figure A2": os.path.join(figures_dir, "privacy_index_construction_diagram_optimized.png"),
    }
    
    # Find figure captions and insert images before them
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        for fig_name, img_path in images.items():
            if text.startswith(fig_name + ":"):
                print(f"   Found {fig_name} caption at paragraph {i}")
                
                # Check if there's already an image (run with drawing)
                # Insert image in the paragraph before this one
                if i > 0:
                    prev_para = doc.paragraphs[i-1]
                    # Clear any placeholder text
                    if "Insert" in prev_para.text or prev_para.text.strip() == "":
                        prev_para.clear()
                    
                    # Add the image
                    run = prev_para.add_run()
                    run.add_picture(img_path, width=Inches(5.5))
                    prev_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    print(f"   ✅ Embedded {fig_name}: {os.path.basename(img_path)}")
    
    print(f"💾 Saving: {output_path}")
    doc.save(output_path)
    print("✅ Done!")

def main():
    input_dir = "/Users/wuyiming/code/thesis/submission"
    
    # Only do identified version as requested
    embed_images(
        os.path.join(input_dir, "Wu_DataSharing_FINAL_PRO.docx"),
        os.path.join(input_dir, "Wu_DataSharing_SUBMIT.docx")
    )

if __name__ == "__main__":
    main()
