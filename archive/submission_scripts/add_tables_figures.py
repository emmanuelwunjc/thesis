#!/usr/bin/env python3
"""
Add tables, figures, and appendix to the submission documents.
Restores all tables and figures from the original thesis.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ============================================================================
# TABLE DATA
# ============================================================================

TABLE_1_DATA = {
    "title": "Table 1: Descriptive Statistics",
    "headers": ["Variable", "Full Sample (N=7,278)", "Analysis Sample (N=2,421)", "Diabetic (n=510)", "Non-Diabetic (n=1,911)"],
    "rows": [
        ["Diabetes Status", "", "", "", ""],
        ["  Has Diabetes", "1,534 (21.1%)", "510 (21.1%)", "510 (100.0%)", "0 (0.0%)"],
        ["  No Diabetes", "5,744 (78.9%)", "1,911 (78.9%)", "0 (0.0%)", "1,911 (100.0%)"],
        ["Privacy Caution Index", "", "", "", ""],
        ["  Mean (SD)", "0.47 (0.09)", "0.47 (0.09)", "0.48 (0.10)", "0.47 (0.09)"],
        ["  Range", "0.23–0.78", "0.23–0.78", "0.23–0.78", "0.23–0.78"],
        ["Data Sharing Willingness", "", "", "", ""],
        ["  Willing to Share", "2,026 (27.8%)", "1,636 (67.6%)", "—", "—"],
        ["  Not Willing", "636 (8.7%)", "785 (32.4%)", "—", "—"],
        ["Age (years)", "", "", "", ""],
        ["  Mean (SD)", "58.3 (15.2)", "58.3 (15.2)", "—", "—"],
        ["Gender", "", "", "", ""],
        ["  Female", "3,792 (52.1%)", "1,260 (52.1%)", "—", "—"],
        ["  Male", "3,486 (47.9%)", "1,161 (47.9%)", "—", "—"],
        ["Health Insurance", "", "", "", ""],
        ["  Has Insurance", "6,208 (85.3%)", "2,064 (85.3%)", "—", "—"],
    ],
    "note": "Note: Privacy Caution Index: 0–1 scale, Cronbach's α = 0.78. Full Sample = HINTS 7 Public Dataset (2022)."
}

TABLE_2_DATA = {
    "title": "Table 2: Main Regression Model Results",
    "subtitle": "Dependent Variable: Data Sharing Willingness (binary 0/1); N = 2,421; R² = 0.174",
    "headers": ["Variable", "Coefficient", "Std. Error", "t-statistic", "p-value", ""],
    "rows": [
        ["Constant", "1.667", "0.074", "22.41", "<0.001", "***"],
        ["Diabetes Status", "0.028", "0.020", "1.40", "0.161", ""],
        ["Privacy Caution Index", "−2.316", "0.108", "−21.49", "<0.001", "***"],
        ["Age", "0.002", "0.001", "5.14", "<0.001", "***"],
        ["Education Level", "−0.015", "0.010", "−1.52", "0.129", ""],
    ],
    "note": "Note: Weighted least squares regression. *** p < 0.001, ** p < 0.01, * p < 0.05, † p < 0.10"
}

TABLE_3_DATA = {
    "title": "Table 3: Interaction Model Results",
    "subtitle": "Dependent Variable: Data Sharing Willingness (binary 0/1); N = 2,421; R² = 0.175",
    "headers": ["Variable", "Coefficient", "Std. Error", "t-statistic", "p-value", ""],
    "rows": [
        ["Constant", "1.720", "0.079", "21.89", "<0.001", "***"],
        ["Diabetes Status", "−0.171", "0.098", "−1.75", "0.081", "†"],
        ["Privacy Caution Index", "−2.441", "0.123", "−19.78", "<0.001", "***"],
        ["Diabetes × Privacy Index", "0.490", "0.236", "2.07", "0.038", "*"],
        ["Age", "0.002", "0.001", "4.99", "<0.001", "***"],
        ["Education Level", "−0.014", "0.010", "−1.47", "0.142", ""],
    ],
    "note": "Note: Weighted least squares regression with interaction term. *** p < 0.001, ** p < 0.01, * p < 0.05, † p < 0.10"
}

# Appendix Tables
TABLE_A1_DATA = {
    "title": "Table A1: Age Group Subgroup Analysis",
    "headers": ["Age Group", "n", "Diabetes Coef.", "Privacy Coef.", "R²"],
    "rows": [
        ["18–35 years", "577", "−0.048 (p=0.466)", "−2.894 (p<0.001)", "0.200"],
        ["36–50 years", "637", "0.056 (p=0.175)", "−2.658 (p<0.001)", "0.192"],
        ["51–65 years", "649", "0.016 (p=0.627)", "−2.366 (p<0.001)", "0.167"],
        ["65+ years", "546", "0.011 (p=0.728)", "−1.640 (p<0.001)", "0.130"],
    ],
    "note": "Note: Separate regressions run for each age group. Privacy coefficient significant across all groups."
}

TABLE_A2_DATA = {
    "title": "Table A2: Privacy Caution Index Sub-dimensions",
    "headers": ["Sub-dimension", "Diabetic Mean (SD)", "Non-Diabetic Mean (SD)", "Difference"],
    "rows": [
        ["Sharing Willingness", "0.49 (0.14)", "0.54 (0.12)", "−0.04"],
        ["Portal Usage", "0.61 (0.24)", "0.62 (0.23)", "−0.01"],
        ["Device Usage", "0.42 (0.28)", "0.34 (0.26)", "+0.08"],
        ["Trust Levels", "0.28 (0.17)", "0.27 (0.16)", "+0.01"],
        ["Social Media", "0.56 (0.10)", "0.54 (0.09)", "+0.02"],
        ["Overall Index", "0.48 (0.10)", "0.47 (0.09)", "+0.01"],
    ],
    "note": "Note: Higher values indicate more privacy caution. Index range: 0–1."
}

# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def set_cell_font(cell, size=10, bold=False):
    """Set font for all runs in a cell."""
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = RGBColor(0, 0, 0)

def add_table_to_doc(doc, table_data, after_para_idx=None):
    """Add a formatted table to the document."""
    
    # Find insertion point
    if after_para_idx is not None:
        # Insert after specific paragraph
        para = doc.paragraphs[after_para_idx]
    else:
        para = doc.paragraphs[-1]
    
    # Add title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(table_data["title"])
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(12)
    title_run.font.bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(12)
    title_para.paragraph_format.space_after = Pt(6)
    
    # Add subtitle if exists
    if "subtitle" in table_data:
        sub_para = doc.add_paragraph()
        sub_run = sub_para.add_run(table_data["subtitle"])
        sub_run.font.name = 'Times New Roman'
        sub_run.font.size = Pt(10)
        sub_run.font.italic = True
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_para.paragraph_format.space_after = Pt(6)
    
    # Create table
    num_cols = len(table_data["headers"])
    num_rows = len(table_data["rows"]) + 1  # +1 for header
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Add headers
    header_row = table.rows[0]
    for i, header in enumerate(table_data["headers"]):
        cell = header_row.cells[i]
        cell.text = header
        set_cell_font(cell, size=10, bold=True)
    
    # Add data rows
    for row_idx, row_data in enumerate(table_data["rows"]):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = str(cell_text)
            # First column left-aligned
            if col_idx == 0:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_cell_font(cell, size=10, bold=False)
    
    # Add note
    note_para = doc.add_paragraph()
    note_run = note_para.add_run(table_data["note"])
    note_run.font.name = 'Times New Roman'
    note_run.font.size = Pt(9)
    note_run.font.italic = True
    note_para.paragraph_format.space_before = Pt(3)
    note_para.paragraph_format.space_after = Pt(12)
    
    return table

def add_figure_placeholder(doc, figure_num, caption, filename):
    """Add a figure placeholder with caption."""
    
    # Add figure number and filename reference
    fig_para = doc.add_paragraph()
    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_para.paragraph_format.space_before = Pt(12)
    
    placeholder = fig_para.add_run(f"[Insert {filename} here]")
    placeholder.font.name = 'Times New Roman'
    placeholder.font.size = Pt(10)
    placeholder.font.italic = True
    placeholder.font.color.rgb = RGBColor(128, 128, 128)
    
    # Add caption
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_para.add_run(f"Figure {figure_num}: {caption}")
    cap_run.font.name = 'Times New Roman'
    cap_run.font.size = Pt(10)
    cap_run.font.bold = True
    cap_para.paragraph_format.space_after = Pt(12)

def find_section_paragraph(doc, section_name):
    """Find the paragraph index for a section heading."""
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == section_name:
            return i
    return None

def insert_after_section(doc, section_name, content_func, *args):
    """Insert content after a section heading and its first paragraph."""
    section_idx = find_section_paragraph(doc, section_name)
    if section_idx is not None:
        # Find end of first paragraph after section heading
        for i in range(section_idx + 1, len(doc.paragraphs)):
            para = doc.paragraphs[i]
            # Skip empty paragraphs
            if para.text.strip():
                # Found first non-empty paragraph, insert after next paragraph break
                return i + 1
    return None

# ============================================================================
# MAIN DOCUMENT PROCESSING
# ============================================================================

def add_tables_and_figures(input_path, output_path):
    """Add all tables and figures to the document."""
    
    print(f"📄 Loading: {input_path}")
    doc = Document(input_path)
    
    # Track where to insert tables/figures
    sections = {}
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text in ["Data and Methods", "Results", "Conceptual Framework", 
                    "Discussion", "Contributions", "Limitations", "References"]:
            sections[text] = i
    
    print(f"   Found sections: {list(sections.keys())}")
    
    # We need to add content at the END of sections, which is tricky
    # Instead, we'll add placeholders that tell the user where to insert
    
    # For now, let's add all tables and figures at the end before References
    ref_idx = sections.get("References", len(doc.paragraphs) - 1)
    
    # Add a separator before tables
    sep = doc.add_paragraph()
    sep.paragraph_format.page_break_before = True
    
    # ========== MAIN BODY TABLES ==========
    print("📊 Adding main body tables...")
    
    # Table 1: Descriptive Statistics
    add_table_to_doc(doc, TABLE_1_DATA)
    
    # Table 2: Main Regression
    add_table_to_doc(doc, TABLE_2_DATA)
    
    # Table 3: Interaction Model
    add_table_to_doc(doc, TABLE_3_DATA)
    
    # ========== FIGURES ==========
    print("🖼️  Adding figure placeholders...")
    
    add_figure_placeholder(doc, 1, 
        "Conceptual Framework: Privacy Caution, Diabetes Status, and Data Sharing",
        "conceptual_framework_diagram.png")
    
    # ========== APPENDIX ==========
    print("📎 Adding appendix...")
    
    # Appendix header
    app_header = doc.add_paragraph()
    app_header.paragraph_format.page_break_before = True
    app_run = app_header.add_run("APPENDIX")
    app_run.font.name = 'Times New Roman'
    app_run.font.size = Pt(14)
    app_run.font.bold = True
    app_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    app_header.paragraph_format.space_after = Pt(24)
    
    # Appendix A: Variable Definitions
    app_a = doc.add_paragraph()
    app_a_run = app_a.add_run("Appendix A: Variable Definitions")
    app_a_run.font.name = 'Times New Roman'
    app_a_run.font.size = Pt(12)
    app_a_run.font.bold = True
    app_a.paragraph_format.space_before = Pt(12)
    app_a.paragraph_format.space_after = Pt(6)
    
    add_figure_placeholder(doc, "A1",
        "Complete Variable Definitions and Coding",
        "variable_definition_table.txt (or create as table)")
    
    # Appendix B: Sample Selection
    app_b = doc.add_paragraph()
    app_b_run = app_b.add_run("Appendix B: Sample Selection Flowchart")
    app_b_run.font.name = 'Times New Roman'
    app_b_run.font.size = Pt(12)
    app_b_run.font.bold = True
    app_b.paragraph_format.space_before = Pt(12)
    app_b.paragraph_format.space_after = Pt(6)
    
    # Add sample selection description
    ss_para = doc.add_paragraph()
    ss_text = """The analysis sample was derived from the HINTS 7 Public Dataset (N=7,278) through the following steps:
(1) Merge with privacy index data; (2) Filter to valid data sharing responses (excluding "Inapplicable" and "Commission Error" responses, removing 4,616 cases); (3) Remove observations with missing key variables. Final analytic sample: N=2,421."""
    ss_run = ss_para.add_run(ss_text)
    ss_run.font.name = 'Times New Roman'
    ss_run.font.size = Pt(11)
    ss_para.paragraph_format.first_line_indent = Inches(0.5)
    ss_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    # Appendix C: Age Subgroup Analysis
    app_c = doc.add_paragraph()
    app_c_run = app_c.add_run("Appendix C: Age Group Subgroup Analysis")
    app_c_run.font.name = 'Times New Roman'
    app_c_run.font.size = Pt(12)
    app_c_run.font.bold = True
    app_c.paragraph_format.space_before = Pt(18)
    app_c.paragraph_format.space_after = Pt(6)
    
    add_table_to_doc(doc, TABLE_A1_DATA)
    
    # Appendix D: Privacy Sub-dimensions
    app_d = doc.add_paragraph()
    app_d_run = app_d.add_run("Appendix D: Privacy Caution Index Sub-dimensions")
    app_d_run.font.name = 'Times New Roman'
    app_d_run.font.size = Pt(12)
    app_d_run.font.bold = True
    app_d.paragraph_format.space_before = Pt(18)
    app_d.paragraph_format.space_after = Pt(6)
    
    add_table_to_doc(doc, TABLE_A2_DATA)
    
    # Save
    print(f"💾 Saving: {output_path}")
    doc.save(output_path)
    print("✅ Done!")

def main():
    input_dir = "/Users/wuyiming/code/thesis/submission"
    
    # Process identified version
    add_tables_and_figures(
        os.path.join(input_dir, "Wu_DataSharing_FINAL.docx"),
        os.path.join(input_dir, "Wu_DataSharing_COMPLETE.docx")
    )
    
    # Process anonymous version
    add_tables_and_figures(
        os.path.join(input_dir, "Wu_DataSharing_FINAL_Anonymous.docx"),
        os.path.join(input_dir, "Wu_DataSharing_COMPLETE_Anonymous.docx")
    )

if __name__ == "__main__":
    main()
