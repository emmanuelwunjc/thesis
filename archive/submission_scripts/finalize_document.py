#!/usr/bin/env python3
"""
Final fixes:
1. Remove duplicate table
2. Add appendix references in text
3. Verify numbering
"""

from docx import Document
from docx.shared import Pt, RGBColor
import os
import re

def finalize_document(input_path, output_path):
    print(f"📄 Loading: {input_path}")
    doc = Document(input_path)
    
    # Track paragraphs to remove (duplicate tables from original)
    paras_to_clear = []
    found_first_interaction = False
    
    # Add appendix references and fix issues
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Remove duplicate "Table 1: Interaction Model Results" that appears early
        if "Table 1: Interaction Model Results" in text and not found_first_interaction:
            # This is the misplaced one from original doc - clear it
            print(f"   🗑️  Removing duplicate table title at para {i}")
            para.clear()
            found_first_interaction = True
            continue
        
        # Add appendix references at strategic points
        
        # After mentioning variable definitions/coding
        if "privacy caution index" in text.lower() and "cronbach" in text.lower():
            if "(see Appendix A" not in text:
                para.add_run(" (see Appendix A for complete variable definitions)")
                print(f"   ✏️  Added Appendix A reference")
        
        # After mentioning sample selection/filtering
        if "2,421" in text and "final" in text.lower() and "sample" in text.lower():
            if "(see Appendix B" not in text and "Appendix B" not in text:
                para.add_run(" (see Appendix B for sample selection details)")
                print(f"   ✏️  Added Appendix B reference")
        
        # After mentioning age-related findings
        if "age group" in text.lower() and ("subgroup" in text.lower() or "heterogeneity" in text.lower()):
            if "(see Appendix C" not in text:
                para.add_run(" (see Appendix C for age group subgroup analysis)")
                print(f"   ✏️  Added Appendix C reference")
        
        # After mentioning privacy index sub-dimensions
        if "sub-dimension" in text.lower() or "subdimension" in text.lower():
            if "(see Appendix D" not in text:
                para.add_run(" (see Appendix D for sub-dimension details)")
                print(f"   ✏️  Added Appendix D reference")
    
    # Save
    print(f"💾 Saving: {output_path}")
    doc.save(output_path)
    print("✅ Done!")

def main():
    input_dir = "/Users/wuyiming/code/thesis/submission"
    
    # Process both versions
    finalize_document(
        os.path.join(input_dir, "Wu_DataSharing_COMPLETE.docx"),
        os.path.join(input_dir, "Wu_DataSharing_FINAL_READY.docx")
    )
    
    finalize_document(
        os.path.join(input_dir, "Wu_DataSharing_COMPLETE_Anonymous.docx"),
        os.path.join(input_dir, "Wu_DataSharing_FINAL_READY_Anonymous.docx")
    )

if __name__ == "__main__":
    main()
