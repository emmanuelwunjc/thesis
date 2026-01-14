#!/usr/bin/env python3
"""
Comprehensive in-text reference addition - more aggressive pattern matching.
"""

from docx import Document
from docx.shared import Pt
import os

def add_all_references(input_path, output_path):
    print(f"📄 Loading: {input_path}")
    doc = Document(input_path)
    
    # Track sections
    current_section = None
    section_starts = {}
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text in ["Abstract", "Introduction and Background", "Literature Review", 
                    "Conceptual Framework", "Data and Methods", "Results", 
                    "Discussion", "Contributions", "Limitations", "References", "APPENDIX"]:
            section_starts[text] = i
            current_section = text
    
    print(f"   Sections found at: {section_starts}")
    
    # Track what we've added
    added = set()
    
    # Process paragraphs
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Skip empty, headings, captions
        if not text or len(text) < 20:
            continue
        if text.startswith("Table ") or text.startswith("Figure ") or text.startswith("Appendix"):
            continue
        
        # Determine section
        section = None
        for sec_name, sec_idx in sorted(section_starts.items(), key=lambda x: x[1], reverse=True):
            if i >= sec_idx:
                section = sec_name
                break
        
        # Skip non-body sections
        if section in ["Abstract", "References", "APPENDIX", None]:
            continue
        
        # === CONCEPTUAL FRAMEWORK ===
        if section == "Conceptual Framework" and "Figure 1" not in added:
            if "four hypotheses" in text.lower():
                new_text = text.replace(
                    "four hypotheses",
                    "four hypotheses (see Figure 1)"
                )
                para.clear()
                run = para.add_run(new_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                added.add("Figure 1")
                print(f"   ✅ Figure 1 at para {i}")
        
        # === DATA AND METHODS ===
        if section == "Data and Methods":
            # Table 1 ref
            if "Table 1" not in added:
                if ("sample" in text.lower() and "publicly" in text.lower()) or \
                   ("descriptive" in text.lower()):
                    new_text = text + " Table 1 presents descriptive statistics for the analytic sample."
                    para.clear()
                    run = para.add_run(new_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    added.add("Table 1")
                    print(f"   ✅ Table 1 at para {i}")
            
            # Appendix E ref
            if "Appendix E" not in added:
                if "robustness" in text.lower() or "linear probability" in text.lower():
                    if "Appendix E" not in text:
                        new_text = text.replace(".", " (see Appendix E).", 1) if "." in text else text + " (see Appendix E)"
                        para.clear()
                        run = para.add_run(new_text)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        added.add("Appendix E")
                        print(f"   ✅ Appendix E at para {i}")
        
        # === RESULTS ===
        if section == "Results":
            # Table 2 ref (main model)
            if "Table 2" not in added:
                if "privacy caution" in text.lower() and "hinder" in text.lower():
                    new_text = text + " Table 2 presents the main regression results."
                    para.clear()
                    run = para.add_run(new_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    added.add("Table 2")
                    print(f"   ✅ Table 2 at para {i}")
            
            # Table 3 ref (interaction model)
            if "Table 3" not in added:
                if "interaction" in text.lower() and "diabetes" in text.lower():
                    if "(Table 3)" not in text:
                        new_text = text + " Table 3 presents the interaction model results."
                        para.clear()
                        run = para.add_run(new_text)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        added.add("Table 3")
                        print(f"   ✅ Table 3 at para {i}")
        
        # === DISCUSSION ===
        if section == "Discussion":
            if "Appendix C" not in added:
                if "age" in text.lower() and ("older" in text.lower() or "younger" in text.lower() or "65" in text):
                    if "Appendix C" not in text:
                        new_text = text + " (see Appendix C for age subgroup results)"
                        para.clear()
                        run = para.add_run(new_text)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        added.add("Appendix C")
                        print(f"   ✅ Appendix C at para {i}")
            
            if "Appendix D" not in added:
                if "sub-dimension" in text.lower() or "subdimension" in text.lower():
                    new_text = text + " (see Appendix D)"
                    para.clear()
                    run = para.add_run(new_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    added.add("Appendix D")
                    print(f"   ✅ Appendix D at para {i}")
    
    print(f"   📊 Total added: {len(added)} - {sorted(added)}")
    
    print(f"💾 Saving: {output_path}")
    doc.save(output_path)
    print("✅ Done!")

def main():
    input_dir = "/Users/wuyiming/code/thesis/submission"
    
    add_all_references(
        os.path.join(input_dir, "Wu_DataSharing_SUBMIT.docx"),
        os.path.join(input_dir, "Wu_DataSharing_FINAL_SUBMIT.docx")
    )

if __name__ == "__main__":
    main()
