#!/usr/bin/env python3
"""
Final fixes v2:
1. Remove duplicate table
2. Add appendix references ONLY in main body (before APPENDIX section)
3. Clean formatting
"""

from docx import Document
from docx.shared import Pt, RGBColor
import os

def finalize_document(input_path, output_path):
    print(f"📄 Loading: {input_path}")
    doc = Document(input_path)
    
    # Find where APPENDIX section starts
    appendix_start = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == "APPENDIX":
            appendix_start = i
            print(f"   📍 Appendix starts at paragraph {i}")
            break
    
    # Process only paragraphs BEFORE appendix
    found_duplicate_table = False
    refs_added = {"A": False, "B": False, "C": False, "D": False}
    
    for i, para in enumerate(doc.paragraphs):
        # Stop at appendix
        if appendix_start and i >= appendix_start:
            break
            
        text = para.text.strip()
        
        # Remove duplicate "Table 1: Interaction Model Results" 
        if "Table 1: Interaction Model Results" in text and not found_duplicate_table:
            print(f"   🗑️  Removing duplicate table at para {i}")
            para.clear()
            found_duplicate_table = True
            continue
        
        # Add appendix references (only once each, in main body)
        
        # Appendix A: After mentioning Cronbach's alpha or index construction
        if not refs_added["A"] and "cronbach" in text.lower():
            para.add_run(" (see Appendix A for complete variable definitions)")
            refs_added["A"] = True
            print(f"   ✏️  Added Appendix A ref at para {i}")
        
        # Appendix B: After mentioning final sample size 2,421
        if not refs_added["B"] and "2,421" in text and ("sample" in text.lower() or "observation" in text.lower()):
            para.add_run(" (see Appendix B for sample selection details)")
            refs_added["B"] = True
            print(f"   ✏️  Added Appendix B ref at para {i}")
        
        # Appendix C: After mentioning age heterogeneity or subgroup
        if not refs_added["C"] and "age" in text.lower() and ("subgroup" in text.lower() or "heterogen" in text.lower() or "older" in text.lower() and "younger" in text.lower()):
            para.add_run(" (see Appendix C)")
            refs_added["C"] = True
            print(f"   ✏️  Added Appendix C ref at para {i}")
        
        # Appendix D: After mentioning sub-dimensions
        if not refs_added["D"] and ("sub-dimension" in text.lower() or "subdimension" in text.lower() or "six component" in text.lower()):
            para.add_run(" (see Appendix D)")
            refs_added["D"] = True
            print(f"   ✏️  Added Appendix D ref at para {i}")
    
    print(f"   📊 Refs added: {refs_added}")
    
    # Save
    print(f"💾 Saving: {output_path}")
    doc.save(output_path)
    print("✅ Done!")

def main():
    input_dir = "/Users/wuyiming/code/thesis/submission"
    
    # Start fresh from COMPLETE versions (before bad refs were added)
    finalize_document(
        os.path.join(input_dir, "Wu_DataSharing_COMPLETE.docx"),
        os.path.join(input_dir, "Wu_DataSharing_READY.docx")
    )
    
    finalize_document(
        os.path.join(input_dir, "Wu_DataSharing_COMPLETE_Anonymous.docx"),
        os.path.join(input_dir, "Wu_DataSharing_READY_Anonymous.docx")
    )

if __name__ == "__main__":
    main()
