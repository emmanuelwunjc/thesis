#!/usr/bin/env python3
"""
Apply academic journal formatting to a Word document.
- 12pt Times New Roman
- Double spacing
- 1-inch margins
- Black text
- First-line indent for body paragraphs
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys
import os

def set_document_margins(doc, margin=1.0):
    """Set all margins to specified inches."""
    for section in doc.sections:
        section.top_margin = Inches(margin)
        section.bottom_margin = Inches(margin)
        section.left_margin = Inches(margin)
        section.right_margin = Inches(margin)

def format_paragraph(para, is_heading=False, is_abstract=False, is_first_body=False):
    """Apply formatting to a paragraph."""
    
    # Set paragraph formatting
    para_format = para.paragraph_format
    
    if is_heading:
        # Headings: bold, no indent, space before
        para_format.space_before = Pt(12)
        para_format.space_after = Pt(6)
        para_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        para_format.first_line_indent = Inches(0)
    elif is_abstract:
        # Abstract: no indent, double spaced
        para_format.first_line_indent = Inches(0)
        para_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        para_format.space_after = Pt(0)
    else:
        # Body paragraphs: 0.5 inch first line indent, double spaced
        para_format.first_line_indent = Inches(0.5)
        para_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        para_format.space_after = Pt(0)
        para_format.space_before = Pt(0)
    
    # Format all runs in paragraph
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        
        # Set font for complex scripts (ensures Times New Roman everywhere)
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:cs'), 'Times New Roman')
        
        if is_heading:
            run.font.bold = True

def format_document(input_path, output_path):
    """Apply academic formatting to the document."""
    
    print(f"📄 Loading: {input_path}")
    doc = Document(input_path)
    
    # Set margins
    print("📐 Setting 1-inch margins...")
    set_document_margins(doc, 1.0)
    
    # Track section for formatting decisions
    in_abstract = False
    in_references = False
    first_body_para = True
    
    print("✍️  Applying formatting to paragraphs...")
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Detect section
        if text == "Abstract":
            in_abstract = True
            format_paragraph(para, is_heading=True)
            continue
        elif text.startswith("Keywords:") or text.startswith("**Keywords"):
            in_abstract = False
            format_paragraph(para, is_abstract=True)  # No indent for keywords
            continue
        elif text in ["Introduction and Background", "Literature Review", 
                      "Conceptual Framework", "Data and Methods", "Results",
                      "Discussion", "Contributions", "Limitations", "References"]:
            in_abstract = False
            if text == "References":
                in_references = True
            format_paragraph(para, is_heading=True)
            first_body_para = True
            continue
        
        # Format based on section
        if in_abstract:
            format_paragraph(para, is_abstract=True)
        elif in_references:
            # References: hanging indent
            para.paragraph_format.first_line_indent = Inches(-0.5)
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)
        else:
            format_paragraph(para, is_first_body=first_body_para)
            first_body_para = False
    
    # Format tables
    print("📊 Formatting tables...")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)  # Slightly smaller for tables
                        run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Save
    print(f"💾 Saving: {output_path}")
    doc.save(output_path)
    print("✅ Done!")

def main():
    input_dir = "/Users/wuyiming/code/thesis/submission"
    
    # Format the identified version
    format_document(
        os.path.join(input_dir, "Wu_DataSharing_Pandoc.docx"),
        os.path.join(input_dir, "Wu_DataSharing_FINAL.docx")
    )
    
    # Format the anonymous version
    format_document(
        os.path.join(input_dir, "Wu_DataSharing_Pandoc_Anonymous.docx"),
        os.path.join(input_dir, "Wu_DataSharing_FINAL_Anonymous.docx")
    )

if __name__ == "__main__":
    main()
