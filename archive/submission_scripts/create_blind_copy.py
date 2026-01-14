#!/usr/bin/env python3
"""
Create blind review copy - remove all identifying information.
"""

from docx import Document
from docx.shared import Pt, Inches
import os

def create_blind_copy(input_path, output_path):
    print(f"📄 Loading: {input_path}")
    doc = Document(input_path)
    
    removals = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Remove author name
        if text == "Yiming Wu":
            para.clear()
            para.add_run("[Author]")
            removals.append("Author name")
            print(f"   🔒 Redacted author name at para {i}")
        
        # Remove institution
        if text == "Georgetown University":
            para.clear()
            para.add_run("[Institution]")
            removals.append("Institution")
            print(f"   🔒 Redacted institution at para {i}")
        
        # Remove email
        if "@georgetown.edu" in text or "@" in text and "georgetown" in text.lower():
            para.clear()
            para.add_run("[Email redacted for blind review]")
            removals.append("Email")
            print(f"   🔒 Redacted email at para {i}")
        
        # Check for any other identifying info in text
        if "Yiming" in text or "Wu, Y" in text:
            new_text = text.replace("Yiming Wu", "[Author]").replace("Yiming", "[Author]").replace("Wu, Y", "[Author]")
            para.clear()
            run = para.add_run(new_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            removals.append("Name in text")
            print(f"   🔒 Redacted name in text at para {i}")
    
    print(f"   📊 Total redactions: {len(removals)}")
    
    print(f"💾 Saving: {output_path}")
    doc.save(output_path)
    print("✅ Done!")

def main():
    input_dir = "/Users/wuyiming/code/thesis/submission"
    
    create_blind_copy(
        os.path.join(input_dir, "Wu_DataSharing_FINAL_SUBMIT.docx"),
        os.path.join(input_dir, "Wu_DataSharing_BLIND_REVIEW.docx")
    )

if __name__ == "__main__":
    main()
